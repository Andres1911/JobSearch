from openai import AsyncOpenAI
import json
import time
from datetime import datetime, timedelta
from typing import Dict, List, Optional
from dataclasses import dataclass
from pathlib import Path
import re
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from personal_context import PERSONAL_CONTEXT, PersonalContext
from prompts_new import COMPANY_RESEARCH_PROMPT, COVER_LETTER_GENERATION_PROMPT_ENGLISH, COVER_LETTER_GENERATION_PROMPT_FRENCH

@dataclass
class JobPosting:
    company_name: str
    job_title: str
    job_description: str
    job_id: str
    link: Optional[str] = None

class CoverLetterGenerator:
    def __init__(self, openai_api_key: str, personal_context: PersonalContext, config_file="user_config.json"):
        self.client = AsyncOpenAI(api_key=openai_api_key)
        self.personal_context = personal_context

        # Load user configuration
        self.config = self.load_config(config_file)

        # New organized structure: ./job_contents/Company_Name/
        self.job_contents_dir = Path("./job_contents")
        self.company_research_cache = Path("./company_research_cache.json")

        # Professional contact information from config
        personal_info = self.config.get('personal_info', {})
        self.contact_info = {
            "name": personal_info.get('name', 'Your Name'),
            "email": personal_info.get('email', 'your.email@example.com'),
            "phone": personal_info.get('phone', '(000) 000-0000'),
            "linkedin": personal_info.get('linkedin', 'linkedin.com/in/yourprofile')
        }

    def load_config(self, config_file):
        """Load configuration from JSON file"""
        try:
            with open(config_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except FileNotFoundError:
            print(f"Config file {config_file} not found. Using default values.")
            return {}
        except json.JSONDecodeError:
            print(f"Error reading config file {config_file}. Using default values.")
            return {}

    def check_existing_cover_letter(self, company_name: str, job_id: str, job_title: str = None) -> bool:
        """Check if a cover letter already exists for this job"""
        company_folder = self.job_contents_dir / company_name
        if not company_folder.exists():
            return False

        # If we have job_title, check for position-specific files
        if job_title:
            safe_title = self.get_safe_filename(job_title)
            pdf_file = company_folder / f"{safe_title}_cover_letter.pdf"
            docx_file = company_folder / f"{safe_title}_cover_letter.docx"
            if pdf_file.exists() or docx_file.exists():
                return True

        # Fallback: check for generic files (backward compatibility)
        pdf_file = company_folder / "cover_letter.pdf"
        docx_file = company_folder / "cover_letter.docx"

        return pdf_file.exists() or docx_file.exists()

    def ensure_company_folder(self, company_name: str) -> Path:
        """Create company folder if it doesn't exist and return path"""
        company_folder = self.job_contents_dir / company_name
        company_folder.mkdir(parents=True, exist_ok=True)
        return company_folder

    def save_job_description(self, company_folder: Path, job_posting: JobPosting):
        """Save job description to company folder"""
        safe_title = self.get_safe_filename(job_posting.job_title)
        job_desc_file = company_folder / f"{safe_title}_job_description.md"
        with open(job_desc_file, 'w', encoding='utf-8') as f:
            f.write(f"# {job_posting.job_title}\n\n")
            f.write(f"**Company:** {job_posting.company_name}\n")
            f.write(f"**Job ID:** {job_posting.job_id}\n")
            if job_posting.link:
                f.write(f"**Link:** {job_posting.link}\n")
            f.write(f"\n## Job Description\n\n{job_posting.job_description}")

    def save_research_results(self, company_folder: Path, company_name: str, research_data: dict):
        """Save research results to company folder"""
        research_file = company_folder / f"Research_{company_name.replace(' ', '_')}.json"
        research_with_timestamp = {
            'timestamp': datetime.now().isoformat(),
            'data': research_data
        }
        with open(research_file, 'w', encoding='utf-8') as f:
            json.dump(research_with_timestamp, f, indent=2)

    @staticmethod
    def get_safe_company_name(company_name: str) -> str:
        """Convert company name to safe folder name"""
        return "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip().replace(' ', '_')

    @staticmethod
    def get_safe_filename(text: str) -> str:
        """Convert text to safe filename by removing/replacing unsafe characters"""
        # Replace spaces with underscores and remove unsafe characters
        safe_text = "".join(c for c in text if c.isalnum() or c in (' ', '-', '_')).strip()
        return safe_text.replace(' ', '_')

    def get_company_folder(self, company_name: str) -> Path:
        """Get the company folder path, create if it doesn't exist"""
        safe_name = self.get_safe_company_name(company_name)
        company_folder = self.job_contents_dir / safe_name
        company_folder.mkdir(parents=True, exist_ok=True)
        return company_folder

    def load_company_research_cache(self) -> Dict:
        """Load cached company research results"""
        if self.company_research_cache.exists():
            try:
                with open(self.company_research_cache, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                print(f"⚠️ Error loading company research cache: {e}")
        return {}

    def save_company_research_cache(self, cache: Dict):
        """Save company research results to cache"""
        try:
            print(f"💾 Saving company research cache to: {self.company_research_cache}")
            with open(self.company_research_cache, 'w', encoding='utf-8') as f:
                json.dump(cache, f, indent=2, ensure_ascii=False)
            print(f"✅ Successfully saved cache with {len(cache)} companies")
        except Exception as e:
            print(f"❌ Error saving company research cache: {e}")
            print(f"   Cache path: {self.company_research_cache}")
            print(f"   Cache data keys: {list(cache.keys()) if cache else 'None'}")
            raise  # Re-raise to see the full error

    @staticmethod
    def is_research_fresh(cached_data: Dict) -> bool:
        """Check if cached research is less than 2 years old"""
        if 'timestamp' not in cached_data:
            return False

        cached_time = datetime.fromisoformat(cached_data['timestamp'])
        return datetime.now() - cached_time < timedelta(days=730) # 2 years

    def load_job_description(self, company_name: str, job_title: str, job_id: str) -> Optional[str]:
        """Load job description from company folder"""
        try:
            company_folder = self.get_company_folder(company_name)
            # Use job title for filename
            safe_title = self.get_safe_filename(job_title)
            job_desc_file = company_folder / f"{safe_title}_job_description.md"

            if not job_desc_file.exists():
                # Fallback: try old naming with job ID
                old_job_id_file = company_folder / f"Job_Description_{job_id}.md"
                if old_job_id_file.exists():
                    job_desc_file = old_job_id_file
                else:
                    # Fallback: try generic filename (for backward compatibility)
                    generic_file = company_folder / "job_description.md"
                    if generic_file.exists():
                        job_desc_file = generic_file
                    else:
                        # Fallback: try old location with job title
                        old_filepath = self.job_contents_dir / f"{safe_title}_{job_id}.md"
                        if old_filepath.exists():
                            print(f"📁 Moving job description to new structure: {company_name}")
                            # Copy to new location
                            with open(old_filepath, 'r', encoding='utf-8') as f:
                                content = f.read()

                            with open(job_desc_file, 'w', encoding='utf-8') as f:
                                f.write(content)

                            print(f"✅ Moved: {old_filepath} -> {job_desc_file}")
                        else:
                            print(f"⚠️ Job description file not found: {job_desc_file}")
                            return None

            with open(job_desc_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Extract job description content
            description_match = re.search(r'## Job Description\s*\n\n(.*)', content, re.DOTALL)
            if description_match:
                return description_match.group(1).strip()

            lines = content.split('\n')
            description_lines = []
            found_description = False
            for line in lines:
                if '---' in line:
                    found_description = True
                    continue
                if found_description:
                    description_lines.append(line)
            return '\n'.join(description_lines).strip()

        except Exception as e:
            print(f"❌ Error loading job description for {job_title} at {company_name}: {e}")
            return None

    async def research_company(self, company_name: str) -> Dict:
        """Research company with caching to save tokens"""
        # Load cache
        cache = self.load_company_research_cache()

        # Check if we have fresh cached data (less than 2 years old)
        if company_name in cache and self.is_research_fresh(cache[company_name]):
            print(f"💰 Using cached research for {company_name} (saves tokens!)")
            return cache[company_name]['data']

        print(f"🔍 Researching {company_name} (using AI tokens)")

        raw_response = ""  # Initialize to avoid reference before assignment warning
        try:
            # Use the imported prompt and format it with the company name
            prompt = COMPANY_RESEARCH_PROMPT.format(company_name=company_name)

            response = await self.client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],  # Type: ignore - OpenAI SDK handles this correctly
                max_tokens=1000
            )

            raw_response = response.choices[0].message.content
            print(f"🔍 Raw AI response length: {len(raw_response)} characters")

            # Strip markdown code blocks if present
            json_content = raw_response.strip()
            if json_content.startswith('```json'):
                json_content = json_content[7:]  # Remove ```json
            if json_content.startswith('```'):
                json_content = json_content[3:]   # Remove ```
            if json_content.endswith('```'):
                json_content = json_content[:-3]  # Remove ending ```
            json_content = json_content.strip()

            research_data = json.loads(json_content)
            print(f"✅ Successfully parsed JSON research data")

            # Cache the results with timestamp
            cache[company_name] = {
                'data': research_data,
                'timestamp': datetime.now().isoformat()
            }
            self.save_company_research_cache(cache)

            # Also save to company folder
            company_folder = self.get_company_folder(company_name)
            self.save_research_results(company_folder, company_name, research_data)

            return research_data

        except json.JSONDecodeError as e:
            print(f"❌ JSON parsing failed: {e}")
            print(f"❌ Raw response that failed to parse: {raw_response}")
            return {"error": "Failed to parse company research"}
        except Exception as e:
            print(f"❌ Error researching {company_name}: {e}")
            return {"error": f"Research failed: {str(e)}"}

    @staticmethod
    def detect_language(text: str) -> str:
        """Detect if text is primarily in French or English"""
        # Common French words and phrases that indicate French content
        french_indicators = [
            # Common French words
            'et', 'de', 'le', 'la', 'les', 'du', 'des', 'un', 'une', 'ce', 'cette', 'nous', 'vous', 'ils', 'elles',
            # French job-specific terms
            'stagiaire', 'stage', 'développement', 'logiciel', 'équipe', 'projet', 'compétences', 'expérience',
            'recherche', 'candidat', 'poste', 'entreprise', 'formation', 'université', 'étudiant',
            # French verbs
            'être', 'avoir', 'faire', 'aller', 'voir', 'savoir', 'pouvoir', 'falloir', 'vouloir', 'venir',
            # French technical terms
            'informatique', 'programmation', 'données', 'système', 'application', 'plateforme',
            # French prepositions and articles
            'dans', 'pour', 'avec', 'sur', 'par', 'sans', 'sous', 'entre', 'chez', 'vers'
        ]

        # Convert to lowercase for comparison
        text_lower = text.lower()

        # Count French indicators
        french_count = sum(1 for word in french_indicators if word in text_lower)

        # Simple heuristic: if we find many French indicators, it's likely French
        # For a typical job description, if we find 10+ French indicators, it's probably French
        return "french" if french_count >= 10 else "english"

    def create_killer_cover_letter_prompt(self, job_posting: JobPosting, company_info: Dict) -> str:
        """Create the prompt for generating a killer cover letter"""
        # Detect language of job description
        language = self.detect_language(job_posting.job_description)

        # Choose appropriate prompt
        if language == "french":
            prompt_template = COVER_LETTER_GENERATION_PROMPT_FRENCH
            print(f"🇫🇷 Detected French job description, using French prompt")
        else:
            prompt_template = COVER_LETTER_GENERATION_PROMPT_ENGLISH
            print(f"🇺🇸 Detected English job description, using English prompt")

        # Format personal context
        personal_context_text = f"""
RESUME: {self.personal_context.resume}

SKILLS: {', '.join(self.personal_context.skills)}

ACHIEVEMENTS: {', '.join(self.personal_context.achievements)}

PERSONAL VALUES: {', '.join(self.personal_context.personal_values)}

ADDITIONAL INFO: {self.personal_context.additional_info}

CONSTRAINTS: {self.personal_context.constraints}
        """.strip()

        # Format company research
        company_research_text = json.dumps(company_info, indent=2) if company_info else "No research data available"

        return prompt_template.format(
            personal_context=personal_context_text,
            company_name=job_posting.company_name,
            job_title=job_posting.job_title,
            job_description=job_posting.job_description,
            company_research=company_research_text,
            user_name=self.contact_info['name']  # Use the name from config
        )

    def create_pdf_cover_letter(self, content: str, job_posting: JobPosting, filepath: Path):
        """Generate a professional PDF cover letter"""
        # Detect language for proper formatting
        language = self.detect_language(job_posting.job_description)
        hiring_manager_text = "Responsable du recrutement" if language == "french" else "Hiring Manager"

        # Format date according to language
        if language == "french":
            import locale
            try:
                locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
                date_formatted = time.strftime('%d %B %Y')
            except locale.Error:
                # Fallback if French locale not available
                months_fr = {
                    'January': 'janvier', 'February': 'février', 'March': 'mars', 'April': 'avril',
                    'May': 'mai', 'June': 'juin', 'July': 'juillet', 'August': 'août',
                    'September': 'septembre', 'October': 'octobre', 'November': 'novembre', 'December': 'décembre'
                }
                date_en = time.strftime('%d %B %Y')
                for en, fr in months_fr.items():
                    date_formatted = date_en.replace(en, fr)
                    if date_formatted != date_en:
                        break
        else:
            date_formatted = time.strftime('%B %d, %Y')

        doc = SimpleDocTemplate(str(filepath), pagesize=letter, topMargin=0.5*inch)
        styles = getSampleStyleSheet()

        # Custom styles with better spacing
        header_style = ParagraphStyle(
            'CustomHeader',
            parent=styles['Normal'],
            fontSize=11,
            alignment=2,  # TA_RIGHT equivalent
            spaceAfter=20
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            alignment=0,  # TA_LEFT equivalent
            spaceAfter=14,  # Increased spacing between paragraphs
            leftIndent=0,
            rightIndent=0,
            leading=14  # Line spacing within paragraphs
        )

        story = []

        # Header with contact info - make email and LinkedIn clickable and underlined
        header_text = f"""
        <b>{self.contact_info['name']}</b><br/>
        <a href="mailto:{self.contact_info['email']}" color="blue"><u>{self.contact_info['email']}</u></a><br/>
        {self.contact_info['phone']}<br/>
        <a href="https://{self.contact_info['linkedin']}" color="blue"><u>{self.contact_info['linkedin']}</u></a>
        """
        story.append(Paragraph(header_text, header_style))
        story.append(Spacer(1, 20))

        # Date and company info
        date_text = f"""
        {date_formatted}<br/><br/>
        <b>{job_posting.company_name}</b><br/>
        {hiring_manager_text}<br/>
        Re: {job_posting.job_title}
        """
        story.append(Paragraph(date_text, body_style))
        story.append(Spacer(1, 20))

        # Cover letter content with proper paragraph spacing
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                formatted_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', para.strip())
                story.append(Paragraph(formatted_para, body_style))

        doc.build(story)

    def create_docx_cover_letter(self, content: str, job_posting: JobPosting, filepath: Path):
        """Generate a professional Word document cover letter"""
        # Detect language for proper formatting
        language = self.detect_language(job_posting.job_description)
        hiring_manager_text = "Responsable du recrutement" if language == "french" else "Hiring Manager"

        # Format date according to language
        if language == "french":
            import locale
            try:
                locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
                date_formatted = time.strftime('%d %B %Y')
            except locale.Error:
                # Fallback if French locale not available
                months_fr = {
                    'January': 'janvier', 'February': 'février', 'March': 'mars', 'April': 'avril',
                    'May': 'mai', 'June': 'juin', 'July': 'juillet', 'August': 'août',
                    'September': 'septembre', 'October': 'octobre', 'November': 'novembre', 'December': 'décembre'
                }
                date_en = time.strftime('%d %B %Y')
                for en, fr in months_fr.items():
                    date_formatted = date_en.replace(en, fr)
                    if date_formatted != date_en:
                        break
        else:
            date_formatted = time.strftime('%B %d, %Y')

        doc = Document()

        # Header with contact info - right aligned
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Add name (bold)
        name_run = header_para.add_run(f"{self.contact_info['name']}\n")
        name_run.bold = True

        # Add email (as plain text since python-docx hyperlinks are complex)
        header_para.add_run(f"{self.contact_info['email']}\n")

        # Add phone
        header_para.add_run(f"{self.contact_info['phone']}\n")

        # Add LinkedIn
        header_para.add_run(f"{self.contact_info['linkedin']}")

        # Add space
        doc.add_paragraph()

        # Date and company info
        date_para = doc.add_paragraph()
        date_para.add_run(f"{date_formatted}\n\n")

        company_run = date_para.add_run(f"{job_posting.company_name}\n")
        company_run.bold = True
        date_para.add_run(f"{hiring_manager_text}\n")
        date_para.add_run(f"Re: {job_posting.job_title}\n")

        # Add space
        doc.add_paragraph()

        # Cover letter content - single line breaks
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.save(str(filepath))

    async def generate_cover_letter(self, job_posting: JobPosting) -> Dict:
        try:
            print(f"🔄 Generating killer cover letter for: {job_posting.job_title} at {job_posting.company_name}")

            # Step 1: Research company thoroughly (with caching!)
            company_info = await self.research_company(job_posting.company_name)

            # Step 2: Create killer cover letter using proven formula
            prompt = self.create_killer_cover_letter_prompt(job_posting, company_info)

            response = await self.client.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}],  # Type: ignore - OpenAI SDK handles this correctly
                max_tokens=1200,
                temperature=0.7
            )

            cover_letter = response.choices[0].message.content

            # Step 3: Save with position-specific filenames
            company_folder = self.get_company_folder(job_posting.company_name)
            safe_title = self.get_safe_filename(job_posting.job_title)

            # PDF version (most professional) - save with job title
            pdf_filepath = company_folder / f"{safe_title}_cover_letter.pdf"
            self.create_pdf_cover_letter(cover_letter, job_posting, pdf_filepath)

            # Word version (for easy editing) - save with job title
            docx_filepath = company_folder / f"{safe_title}_cover_letter.docx"
            self.create_docx_cover_letter(cover_letter, job_posting, docx_filepath)

            # Also save job description
            self.save_job_description(company_folder, job_posting)

            print(f"✅ Cover letter saved for specific position:")
            print(f"   📁 Company folder: {company_folder}")
            print(f"   📄 PDF: {pdf_filepath}")
            print(f"   📝 Word: {docx_filepath}")
            print(f"   🔍 Research: {company_folder / f'Research_{job_posting.company_name.replace(' ', '_')}.json'}")
            print(f"   📋 Job desc: {company_folder / f'{safe_title}_job_description.md'}")

            return {
                "success": True,
                "cover_letter": cover_letter,
                "pdf_filepath": str(pdf_filepath),
                "docx_filepath": str(docx_filepath),
                "company_folder": str(company_folder),
                "company_info": company_info,
                "job_id": job_posting.job_id
            }

        except Exception as e:
            print(f"❌ Error generating cover letter for {job_posting.job_title}: {e}")
            return {
                "success": False,
                "error": str(e),
                "job_id": job_posting.job_id,
                "company": job_posting.company_name,
                "title": job_posting.job_title
            }

    async def process_job_batch(self, job_postings: List[JobPosting], delay: int = 2) -> List[Dict]:
        results = []
        print(f"🚀 Starting batch processing for {len(job_postings)} jobs...")

        for i, job_posting in enumerate(job_postings):
            print(f"\n📝 Processing {i+1}/{len(job_postings)}: {job_posting.company_name} - {job_posting.job_title}")

            result = await self.generate_cover_letter(job_posting)
            results.append(result)

            if i < len(job_postings) - 1:
                print(f"⏳ Waiting {delay} seconds...")
                time.sleep(delay)

        return results

    @classmethod
    def load_jobs_from_processed_file(cls, processed_jobs_file: str = "processed_jobs.json") -> List[JobPosting]:
        job_postings = []

        try:
            with open(processed_jobs_file, 'r', encoding='utf-8') as f:
                processed_jobs = json.load(f)

            print(f"📂 Loading {len(processed_jobs)} jobs from {processed_jobs_file}")

            # Use a dummy PersonalContext for loading job descriptions only
            generator_temp = cls("dummy_key", PERSONAL_CONTEXT)

            for job_id, job_info in processed_jobs.items():
                job_description = generator_temp.load_job_description(job_info["company"], job_info["job_title"], job_id)

                if job_description:
                    job_posting = JobPosting(
                        company_name=job_info["company"],
                        job_title=job_info["job_title"],
                        job_description=job_description,
                        job_id=job_id,
                        link=job_info.get("link")
                    )
                    job_postings.append(job_posting)
                    print(f"✅ Loaded: {job_info['job_title']} at {job_info['company']}")
                else:
                    print(f"⚠️ Skipped: {job_info['job_title']} at {job_info['company']}")

            print(f"✅ Successfully loaded {len(job_postings)} jobs with descriptions")
            return job_postings

        except Exception as e:
            print(f"❌ Error loading jobs: {e}")
            return []
